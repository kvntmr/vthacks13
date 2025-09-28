"""
DOC/DOCX Parser Service
Extracts text and metadata from Microsoft Word documents
"""

import os
from typing import Dict, Any, List
from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.table import Table

class DocParser:
    """Parser for DOC and DOCX files"""
    
    def __init__(self):
        """Initialize the DOC parser"""
        pass
    
    async def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOC/DOCX file and extract text content
        
        Args:
            file_path: Path to the DOC/DOCX file
            
        Returns:
            Dictionary containing parsed content
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOC/DOCX file not found: {file_path}")
            
            result = {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_type": "docx",
                "paragraphs": [],
                "tables": [],
                "extracted_text": "",
                "metadata": {},
                "processing_summary": {
                    "total_paragraphs": 0,
                    "total_tables": 0,
                    "total_text_length": 0
                }
            }
            
            # Open the document
            doc = Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            result["metadata"] = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "version": core_props.version or ""
            }
            
            # Extract paragraphs
            all_text = []
            paragraph_count = 0
            
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraph_data = {
                        "paragraph_number": paragraph_count + 1,
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal",
                        "text_length": len(para.text.strip())
                    }
                    result["paragraphs"].append(paragraph_data)
                    all_text.append(para.text.strip())
                    paragraph_count += 1
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                table_data = {
                    "table_number": table_count + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "data": []
                }
                
                # Extract table data
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data["data"].append(row_data)
                
                result["tables"].append(table_data)
                table_count += 1
                
                # Add table text to extracted text
                table_text = "\n".join([" | ".join(row) for row in table_data["data"]])
                all_text.append(f"\n[TABLE {table_count}]\n{table_text}\n")
            
            # Combine all text
            result["extracted_text"] = "\n\n".join(all_text)
            result["processing_summary"] = {
                "total_paragraphs": paragraph_count,
                "total_tables": table_count,
                "total_text_length": len(result["extracted_text"])
            }
            
            return result
            
        except Exception as e:
            return {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_type": "docx",
                "error": f"Failed to parse DOC/DOCX file: {str(e)}",
                "extracted_text": "",
                "paragraphs": [],
                "tables": [],
                "metadata": {},
                "processing_summary": {
                    "total_paragraphs": 0,
                    "total_tables": 0,
                    "total_text_length": 0
                }
            }
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return [".docx", ".doc"]
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported"""
        file_ext = os.path.splitext(file_path.lower())[1]
        return file_ext in self.get_supported_formats()

